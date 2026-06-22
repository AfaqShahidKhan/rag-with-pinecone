"""
src/ingestion/loader.py

Document loading pipeline. Supports PDF and DOCX formats.
Returns a flat list of Document objects with source metadata attached.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from pypdf import PdfReader
from docx import Document as DocxDocument

from src.utils.logger import get_logger

log = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


@dataclass
class Document:
    """
    Minimal document unit — a single page/section of text with provenance metadata.
    Passed through the entire pipeline: load → chunk → embed → upsert → retrieve.
    """
    page_content: str
    metadata: dict = field(default_factory=dict)


# ── PDF ───────────────────────────────────────────────────────────────────────

def load_pdf(path: Path) -> list[Document]:
    """Load a single PDF, returning one Document per page."""
    reader = PdfReader(str(path))
    total = len(reader.pages)
    docs: list[Document] = []

    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()

        if not text:
            log.warning(f"Page {i + 1}/{total} of '{path.name}' yielded no text — skipping.")
            continue

        docs.append(Document(
            page_content=text,
            metadata={
                "source": path.name,
                "source_path": str(path),
                "page": i + 1,
                "total_pages": total,
                "file_type": "pdf",
            }
        ))

    log.info(f"Loaded '{path.name}': {len(docs)}/{total} pages with text.")
    return docs


# ── DOCX ──────────────────────────────────────────────────────────────────────

def load_docx(path: Path) -> list[Document]:
    """
    Load a single DOCX file.
    Paragraphs are grouped into page-sized chunks of ~3000 chars each
    to approximate page-level provenance (DOCX has no native page breaks).
    """
    doc = DocxDocument(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    if not paragraphs:
        log.warning(f"'{path.name}' yielded no text — skipping.")
        return []

    # Group paragraphs into pseudo-pages (~3000 chars each)
    PAGE_SIZE = 3000
    pages: list[str] = []
    current: list[str] = []
    current_len = 0

    for para in paragraphs:
        if current_len + len(para) > PAGE_SIZE and current:
            pages.append("\n\n".join(current))
            current = []
            current_len = 0
        current.append(para)
        current_len += len(para)

    if current:
        pages.append("\n\n".join(current))

    docs = [
        Document(
            page_content=text,
            metadata={
                "source": path.name,
                "source_path": str(path),
                "page": i + 1,
                "total_pages": len(pages),
                "file_type": "docx",
            }
        )
        for i, text in enumerate(pages)
    ]

    log.info(f"Loaded '{path.name}': {len(docs)} sections from {len(paragraphs)} paragraphs.")
    return docs


# ── Directory loader ──────────────────────────────────────────────────────────

def load_documents_from_dir(directory: Path) -> list[Document]:
    """Load all supported documents (PDF, DOCX) from a directory recursively."""
    files = sorted([
        p for p in directory.rglob("*")
        if p.suffix.lower() in SUPPORTED_EXTENSIONS
    ])

    if not files:
        raise FileNotFoundError(
            f"No supported files (.pdf, .docx) found in '{directory}'."
        )

    log.info(f"Found {len(files)} file(s) in '{directory}': "
             f"{[f.name for f in files]}")

    all_docs: list[Document] = []
    for file_path in files:
        if file_path.suffix.lower() == ".pdf":
            all_docs.extend(load_pdf(file_path))
        elif file_path.suffix.lower() == ".docx":
            all_docs.extend(load_docx(file_path))

    log.info(f"Total sections loaded: {len(all_docs)}")
    return all_docs


# ── Backward-compatible alias ─────────────────────────────────────────────────
# Keeps validate_loader.py and any other callers that used the old name working.

load_pdfs_from_dir = load_documents_from_dir